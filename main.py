import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from io import BytesIO
from datetime import datetime

# --- 頁面配置 ---
st.set_page_config(page_title="工地雜支管理系統", layout="centered")

st.markdown("""
    <style>
    div.stButton > button:first-child { width: 100%; height: 3.5em; font-size: 18px; font-weight: bold; }
    .total-preview { 
        background-color: #ffffff; padding: 20px; border-radius: 15px; 
        text-align: center; border: 2px solid #1e88e5; margin-bottom: 25px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }
    </style>
    """, unsafe_allow_html=True)

# --- 初始化 Session State ---
if 'data_list' not in st.session_state:
    st.session_state.data_list = []
if 'location_options' not in st.session_state:
    st.session_state.location_options = ["工務所", "建案 A", "建案 B"]

# --- 輔助函式：排序與自動代號生成 ---
def process_data_and_mapping(data):
    """依照日期排序資料，並根據排序後的工地出現順序生成 A-Z 代號"""
    # 1. 排序
    def sort_key(item):
        try: return datetime.strptime(f"{datetime.now().year}/{item['日期']}", "%Y/%m/%d")
        except: return datetime.max
    sorted_data = sorted(data, key=sort_key)
    
    # 2. 生成代號字典 (Mapping)
    unique_locations = []
    for d in sorted_data:
        if d["工地"] not in unique_locations:
            unique_locations.append(d["工地"])
    
    # chr(65) 是 'A'，依序往後推
    mapping = {loc: chr(65 + i) for i, loc in enumerate(unique_locations)}
    return sorted_data, mapping

# --- 側邊欄：管理選單 ---
with st.sidebar:
    st.header("⚙️ 選項設定")
    new_loc = st.text_input("新增常用工地")
    if st.button("➕ 新增"):
        if new_loc and new_loc not in st.session_state.location_options:
            st.session_state.location_options.append(new_loc)
            st.rerun()
    
    st.divider()
    del_loc = st.selectbox("刪除常用工地", options=st.session_state.location_options)
    if st.button("🗑️ 刪除"):
        if del_loc in st.session_state.location_options:
            st.session_state.location_options.remove(del_loc)
            st.rerun()

# --- 主頁面：總金額預覽 ---
st.title("📂 雜支明細自動化")

if st.session_state.data_list:
    total_amt = sum(d['金額'] for d in st.session_state.data_list)
    st.markdown(f"""
        <div class="total-preview">
            <p style="margin:0; color:#666;">目前累計總餘額</p>
            <h1 style="margin:0; color:{'#d32f2f' if total_amt < 0 else '#1e88e5'};">NT$ {total_amt:,}</h1>
        </div>
    """, unsafe_allow_html=True)

# --- 資料輸入區 ---
with st.expander("🖋️ 快速記帳", expanded=True):
    date_val = st.text_input("日期", value=datetime.now().strftime("%m/%d"))
    content_val = st.text_input("項目內容", placeholder="如：五金、餐費")
    
    col1, col2 = st.columns(2)
    with col1:
        # 直接輸入正數，邏輯會轉負數 (支出)
        raw_amt = st.number_input("金額 (輸入 100 即為支出 100)", step=10, value=0)
    with col2:
        loc_choice = st.selectbox("選擇工地", options=st.session_state.location_options + ["+ 手動輸入"])
        if loc_choice == "+ 手動輸入":
            final_loc = st.text_input("輸入新工地全名")
        else:
            final_loc = loc_choice

    if st.button("🚀 新增至清單"):
        if date_val and content_val and final_loc:
            # 自動轉負數邏輯：支出預設為負數
            actual_amt = -abs(raw_amt) if raw_amt > 0 else raw_amt
            st.session_state.data_list.append({
                "日期": date_val, "內容": content_val, "金額": actual_amt, "工地": final_loc
            })
            st.rerun()

# --- 資料預覽與報表生成 ---
if st.session_state.data_list:
    # 核心邏輯：取得排序後的資料與自動生成的 A-Z 對照表
    sorted_list, loc_mapping = process_data_and_mapping(st.session_state.data_list)
    
    st.subheader("📊 資料預覽 (已自動編號)")
    # 預覽表中直接顯示代號，方便核對
    preview_df = pd.DataFrame([{
        "日期": d["日期"], "項目": d["內容"], "金額": d["金額"], 
        "代號": loc_mapping[d["工地"]], "工地全名": d["工地"]
    } for d in sorted_list])
    st.table(preview_df)

    # --- Word 生成邏輯 ---
    def export_word(data, mapping):
        doc = Document()
        # 設定邊距
        for s in doc.sections:
            s.top_margin = s.bottom_margin = Mm(15)
            s.left_margin = s.right_margin = Mm(15)

        # 標題
        title = doc.add_paragraph("雜支明細表")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(18)
        title.runs[0].bold = True
        
        doc.add_paragraph(f"報告日期：{datetime.now().strftime('%Y/%m/%d')}")
        doc.add_paragraph(f"經手人：_________________")

        # 垂直排列計算 (左側到底再右側)
        rows_per_page = 28
        left_side = data[:rows_per_page]
        right_side = data[rows_per_page:rows_per_page*2]

        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'
        
        # 表頭文字置中與加粗
        headers = ["日期", "項目內容", "金額", "工地"] * 2
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # 填入內容
        last_dl, last_dr = None, None
        for i in range(len(left_side)):
            row = table.add_row().cells
            
            # 左側
            d_l = left_side[i]
            txt_date_l = "" if d_l["日期"] == last_dl else d_l["日期"]
            last_dl = d_l["日期"]
            l_vals = [txt_date_l, d_l["內容"], f"{d_l['金額']:,}", mapping[d_l["工地"]]]
            
            # 右側
            r_vals = [""] * 4
            if i < len(right_side):
                d_r = right_side[i]
                txt_date_r = "" if d_r["日期"] == last_dr else d_r["日期"]
                last_dr = d_r["日期"]
                r_vals = [txt_date_r, d_r["內容"], f"{d_r['金額']:,}", mapping[d_r["工地"]]]

            for idx, val in enumerate(l_vals + r_vals):
                row[idx].text = str(val)
                row[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # 總計
        total = sum(d['金額'] for d in data)
        doc.add_paragraph(f"\n總計金額：NT$ {total:,} 元").alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # 自動生成代號索引表
        doc.add_paragraph("-" * 20)
        doc.add_paragraph("【工地代號對照索引】").bold = True
        for name, code in mapping.items():
            doc.add_paragraph(f"{code} : {name}")

        out = BytesIO()
        doc.save(out)
        out.seek(0)
        return out

    col_del, col_dl = st.columns(2)
    with col_del:
        if st.button("🗑️ 清空所有資料"):
            st.session_state.data_list = []
            st.rerun()
    with col_dl:
        word_file = export_word(sorted_list, loc_mapping)
        st.download_button(
            label="📥 下載 Word 報表",
            data=word_file,
            file_name=f"雜支明細_{datetime.now().strftime('%m%d')}.docx"
        )
